"""DocxConverter — extract text from a .docx file and convert it to KB YAML
via an LLM call so it can be dropped into ai/knowledge_base/increments/.

The LLM is given the extracted document text and asked to emit a YAML
structure that matches the existing KB increment format (feature, domain,
description, endpoints/ui_changes, scenarios).  The caller can then review
the YAML before committing it to the KB.
"""
from __future__ import annotations

from pathlib import Path

from ai.prompts.prompt_templates import PromptTemplate

_DOCX_TO_KB_PROMPT = PromptTemplate("""
You are a test automation architect. Convert the document below into a
SentinelFlux Knowledge Base increment YAML file.

The output MUST be valid YAML that follows this schema exactly:

feature: <snake_case_feature_name>
domain: api  # api | web | mobile | security | a11y
description: |
  <One paragraph describing what this feature does and what needs to be tested.>

# For API features — list every endpoint you can identify:
endpoints:
  - name: <Human readable name>
    path: /path/to/endpoint
    method: GET   # GET | POST | PUT | PATCH | DELETE
    description: What this endpoint does
    auth_required: true
    request_body: <describe fields if documented>
    response_codes: [200, 400, 401, 404]

# For web/UI features — describe components changed:
# ui_changes:
#   - component: Form or page name
#     action: What changed

# Test scenarios — derive from the document:
scenarios:
  - name: <scenario name>
    type: happy_path   # happy_path | error | edge_case | security | a11y
    description: <what to test>

RULES:
- Output ONLY the YAML. No markdown fences, no explanations.
- If the document describes both API and web aspects, pick the dominant domain.
- If you cannot determine a value, omit the field rather than guessing.
- Use snake_case for the feature name.
- Keep descriptions concise (1-2 sentences each).

--- DOCUMENT START ---
{document_text}
--- DOCUMENT END ---
""")


class DocxConverter:

    def __init__(self, ai_client=None):
        self._client = ai_client

    def extract_text(self, path: Path) -> str:
        """Extract all paragraph and table text from a .docx file."""
        try:
            from docx import Document  # python-docx
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required for DOCX conversion. "
                "Install it: pip install python-docx"
            ) from exc

        doc = Document(str(path))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = (para.style.name if para.style else "") or ""
                if style_name.startswith("Heading"):
                    level = style_name.replace("Heading ", "").strip()
                    prefix = "#" * int(level) if level.isdigit() else "##"
                    parts.append(f"{prefix} {text}")
                else:
                    parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    def convert_to_kb_yaml(self, document_text: str) -> str:
        """Use the configured LLM to convert extracted text into KB YAML."""
        if self._client is None:
            raise RuntimeError(
                "AI client required for DOCX→YAML conversion. "
                "Configure an AI client (Mistral or Ollama) in settings."
            )
        prompt = _DOCX_TO_KB_PROMPT.format(document_text=document_text[:6000])
        result = self._client.generate(prompt, max_tokens=2000, temperature=0.1).strip()
        # Strip any accidental markdown fences the model adds
        if result.startswith("```"):
            result = "\n".join(
                line for line in result.splitlines()
                if not line.strip().startswith("```")
            ).strip()
        return result

    def convert_file(self, path: Path) -> str:
        """Full pipeline: read .docx → extract text → LLM → KB YAML string."""
        text = self.extract_text(path)
        return self.convert_to_kb_yaml(text)
